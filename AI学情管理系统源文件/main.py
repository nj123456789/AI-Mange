import sys, json, requests
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import QFont, QColor
from docx import Document

# ========= 数据 =========
SUBJECTS = ["语文", "数学", "英语", "物理", "化学", "生物"]

def load_data(file):
    try:
        return json.load(open(file, "r", encoding="utf-8"))
    except:
        return {}

def save_data(file, data):
    json.dump(data, open(file, "w", encoding="utf-8"), ensure_ascii=False, indent=2)

users = load_data("users.json")
students = load_data("students.json")

# ========= AI =========
API_KEY = "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"

def call_ai(prompt):
    """
    使用 OpenRouter API 调用 AI
    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "openai/gpt-4o-mini",
        "messages": [{"role": "user", "content": prompt}],
        "stream": False
    }
    try:
        r = requests.post(url, headers=headers, json=data, timeout=20)
        if r.status_code != 200:
            return f"AI调用异常：Error code {r.status_code} - {r.text}"
        res_json = r.json()
        return res_json["choices"][0]["message"]["content"]
    except Exception as e:
        return f"AI调用异常：{e}"

# ========= 登录界面 =========
class Login(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("登录系统")
        self.resize(300, 200)
        layout = QVBoxLayout()
        self.u = QLineEdit()
        self.u.setPlaceholderText("用户名")
        self.p = QLineEdit()
        self.p.setPlaceholderText("密码")
        self.p.setEchoMode(QLineEdit.Password)
        self.msg = QLabel("")
        btn_login = QPushButton("登录")
        btn_reg = QPushButton("注册")
        btn_login.clicked.connect(self.login)
        btn_reg.clicked.connect(self.reg)
        layout.addWidget(self.u)
        layout.addWidget(self.p)
        layout.addWidget(btn_login)
        layout.addWidget(btn_reg)
        layout.addWidget(self.msg)
        self.setLayout(layout)

    def login(self):
        if self.u.text() in users and users[self.u.text()] == self.p.text():
            self.main = Main()
            self.main.show()
            self.close()
        else:
            self.msg.setText("登录失败")

    def reg(self):
        users[self.u.text()] = self.p.text()
        save_data("users.json", users)
        self.msg.setText("注册成功")

# ========= 主界面 =========
class Main(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI学情分析系统")
        self.resize(1200, 700)
        layout = QVBoxLayout()
        self.search = QLineEdit()
        self.search.setPlaceholderText("搜索学生...")
        self.search.textChanged.connect(self.refresh)
        layout.addWidget(self.search)
        self.table = QTableWidget()
        self.table.setColumnCount(len(SUBJECTS) + 2)
        self.table.setHorizontalHeaderLabels(["姓名"] + SUBJECTS + ["平均分"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setSortingEnabled(True)
        layout.addWidget(self.table)
        self.ai_output = QTextEdit()
        layout.addWidget(self.ai_output)
        btn_layout = QHBoxLayout()
        btn_add = QPushButton("新增")
        btn_add.clicked.connect(self.add_student)
        btn_delete = QPushButton("删除")
        btn_delete.clicked.connect(self.delete_student)
        btn_ai = QPushButton("AI分析")
        btn_ai.clicked.connect(self.ai_analyze)
        btn_export = QPushButton("导出Word")
        btn_export.clicked.connect(self.export)
        btn_layout.addWidget(btn_add)
        btn_layout.addWidget(btn_delete)
        btn_layout.addWidget(btn_ai)
        btn_layout.addWidget(btn_export)
        layout.addLayout(btn_layout)
        self.setLayout(layout)
        self.refresh()

    def closeEvent(self, event):
        QApplication.quit()

    def refresh(self):
        keyword = self.search.text()
        data = {k: v for k, v in students.items() if keyword in k}
        self.table.setRowCount(len(data))
        for row, (name, scores) in enumerate(data.items()):
            avg = sum(scores.values()) / len(scores)
            self.table.setItem(row, 0, QTableWidgetItem(name))
            for col, sub in enumerate(SUBJECTS, start=1):
                val = scores[sub]
                item = QTableWidgetItem(str(val))
                if val >= 90:
                    item.setBackground(QColor("#16a34a"))
                elif val >= 60:
                    item.setBackground(QColor("#eab308"))
                else:
                    item.setBackground(QColor("#dc2626"))
                self.table.setItem(row, col, item)
            self.table.setItem(row, len(SUBJECTS)+1, QTableWidgetItem(f"{avg:.2f}"))

    def add_student(self):
        name, ok = QInputDialog.getText(self, "新增", "姓名")
        if not ok or not name:
            return
        scores = {}
        for sub in SUBJECTS:
            val, ok = QInputDialog.getInt(self, "成绩", sub, 60, 0, 100)
            if not ok:
                return
            scores[sub] = val
        students[name] = scores
        save_data("students.json", students)
        self.refresh()

    def delete_student(self):
        row = self.table.currentRow()
        if row < 0:
            return
        name = self.table.item(row, 0).text()
        del students[name]
        save_data("students.json", students)
        self.refresh()

    def ai_analyze(self):
        row = self.table.currentRow()
        if row < 0:
            return
        name = self.table.item(row, 0).text()
        scores = students[name]
        prompt = f"{name}成绩{scores}请分析"
        self.ai_output.setText("分析中...")
        QApplication.processEvents()
        result = call_ai(prompt)
        self.ai_output.setText(result)

    def export(self):
        path, _ = QFileDialog.getSaveFileName(self, "保存", "报告.docx", "Word (*.docx)")
        if not path:
            return

        doc = Document()
        doc.add_heading("学情报告", 0)

        for name, scores in students.items():
            doc.add_heading(name, 1)
            for k, v in scores.items():
                doc.add_paragraph(f"{k}: {v}")

            # 添加 AI 分析结果
            # 如果当前选中学生的分析已经生成，则使用，否则可以自动生成一次
            prompt = f"{name}成绩{scores}请分析"
            ai_result = call_ai(prompt)
            doc.add_heading("AI分析", level=2)
            doc.add_paragraph(ai_result)

        doc.save(path)
        QMessageBox.information(self, "成功", "已导出包含AI分析的学情报告")

# ========= 启动 =========
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setFont(QFont("Microsoft YaHei"))
    login = Login()
    login.show()
    sys.exit(app.exec_())